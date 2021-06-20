import docx

doc = docx.Document()
doc.add_paragraph('Hello World!', 'Title')
outputFile = '/Users/tony/Tony/PDF/helloworld.docx'
paraObj1 = doc.add_paragraph('This a second paragraphs.')
paraObj2 = doc.add_paragraph('This a yet another paragraph.')
paraObj1.add_run(' This text is being added to the second paragraph.')
print("Save to file:", outputFile)
doc.save(outputFile)
