import docx

import readDocx

filename = '/Users/tony/Downloads/automate_online-materials/demo.docx'
doc = docx.Document(filename)
print(len(doc.paragraphs))
print(doc.paragraphs[0].text)
print(doc.paragraphs[1].text)
print(len(doc.paragraphs[1].runs))
print(doc.paragraphs[1].runs[0].text)
print(doc.paragraphs[1].runs[1].text)
print(doc.paragraphs[1].runs[2].text)
print(doc.paragraphs[1].runs[3].text)
print(doc.paragraphs[1].runs[4].text)

print('----------split line---------------')
print(readDocx.getText(filename))

print(doc.paragraphs[0].text)
print(doc.paragraphs[0].style)
doc.paragraphs[0].style = 'Normal'
print(doc.paragraphs[1].text)
print((doc.paragraphs[1].runs[0].text, doc.paragraphs[1].runs[1].text, doc.paragraphs[1].runs[2].text,
       doc.paragraphs[1].runs[3].text))
doc.paragraphs[1].runs[0].style = 'Quote Char'
doc.paragraphs[1].runs[1].underline = True
doc.paragraphs[1].runs[3].underline = True
outputFile = '/Users/tony/Tony/PDF/restyled.docx'
doc.save(outputFile)
