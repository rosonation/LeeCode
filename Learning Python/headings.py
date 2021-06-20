import docx

doc = docx.Document()
doc.add_heading('herder 0', 0)
doc.add_heading('herder 1', 1)
doc.add_heading('herder 2', 2)
doc.add_heading('herder 3', 3)
doc.add_heading('herder 4', 4)
outputFile = '/Users/tony/Tony/PDF/headings.docx'
print("Save docx to file:", outputFile)
doc.save(outputFile)
